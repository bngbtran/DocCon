import io
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

from .utils import crop_image, html_to_table, pil_to_bytes

_TEXT_LABELS = {
    "text",
    "other_text",
    "paragraph",
    "list",
    "footer",
    "header",
    "caption",
}
_HEADING_LABELS = {
    "title": 1,
    "section_title": 1,
    "paragraph_title": 2,
    "sub_paragraph_title": 3,
}
_FIGURE_LABELS = {"figure", "image"}

_ALIGN_MAP = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
}


class DocxBuilder:
    def __init__(self):
        self.doc = Document()
        self._margins_set = False
        self._page_info: dict = {}
        for p in self.doc.paragraphs:
            p._element.getparent().remove(p._element)

    def set_page_margins(self, page_info: dict):
        if self._margins_set:
            return
        self._page_info = page_info
        sec = self.doc.sections[0]

        def _emu(pt):
            return int(pt * 12700)

        sec.left_margin = _emu(page_info.get("left_margin_pt", 89.9))
        sec.right_margin = _emu(page_info.get("right_margin_pt", 89.1))
        sec.top_margin = _emu(page_info.get("top_margin_pt", 72.0))
        sec.bottom_margin = _emu(page_info.get("bottom_margin_pt", 72.0))
        sec.page_width = _emu(page_info.get("page_width_pt", 595.0))
        sec.page_height = _emu(page_info.get("page_height_pt", 842.0))
        self._margins_set = True

    def add_page_break(self):
        self.doc.add_page_break()

    def process_blocks(self, items: List[dict], page_image: Optional[Image.Image]):
        for item in items:
            itype = item.get("item_type")
            if itype == "columns":
                self._render_columns(item)
            elif itype == "hline":
                self._render_hline(item)
            elif itype in ("paragraph", None) and "alignment" in item:
                self._render_rich(item)
            else:
                self._render_plain(item, page_image)

    def save(self, output_path: str):
        self.doc.save(output_path)

    def _render_rich(self, block: dict, para=None):
        spans = block.get("block_spans", [])
        alignment = block.get("alignment", "left")
        has_bullet = block.get("has_bullet", False)
        indent_left = block.get("indent_left_pt", 0.0)
        indent_hanging = block.get("indent_hanging_pt", 0.0)
        indent_first = block.get("indent_first_line_pt", 0.0)
        space_after = block.get("space_after_pt", 14.0)

        if para is None:
            para = self.doc.add_paragraph()
        pf = para.paragraph_format
        pf.alignment = _ALIGN_MAP.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        pf.space_before = Pt(0)
        pf.space_after = Pt(space_after)

        if has_bullet:
            pf.left_indent = Pt(indent_left)
            pf.first_line_indent = Pt(-indent_hanging)
            self._set_tab_stop(para, indent_left)
            body_font = spans[0]["font"] if spans else "Times New Roman"
            body_size = spans[0]["size"] if spans else 12
            br = para.add_run("•\t")
            br.font.name = body_font
            br.font.size = Pt(body_size)
        elif indent_first > 0:
            pf.first_line_indent = Pt(indent_first)

        _apply_spans(para, spans)

    def _render_columns(self, item: dict):
        cells_data: List[List[dict]] = item.get("cells", [])
        n = len(cells_data)
        if n == 0:
            return

        page_w = item.get("page_width_pt", 595.0)
        left_m = item.get("left_margin_pt", 72.0)
        right_m = item.get("right_margin_pt", 72.0)
        left_edge = left_m
        right_edge = page_w - right_m

        cell_x1s = [min(b["block_bbox"][0] for b in cell) for cell in cells_data]
        cell_x2s = [max(b["block_bbox"][2] for b in cell) for cell in cells_data]

        actual_right = max(cell_x2s[-1], right_edge)

        boundaries = [left_edge]
        for i in range(n - 1):
            mid = (cell_x2s[i] + cell_x1s[i + 1]) / 2.0
            boundaries.append(mid)
        boundaries.append(actual_right)

        col_w_pts = [max(1.0, boundaries[i + 1] - boundaries[i]) for i in range(n)]

        table = self.doc.add_table(rows=1, cols=n)
        table.style = "Table Grid"

        table_w_twips = int((actual_right - left_edge) * 20)
        tbl = table._tbl
        tblPr = _get_or_add(tbl, "w:tblPr")
        _set_child(tblPr, "w:tblW", {"w:w": str(table_w_twips), "w:type": "dxa"})
        _set_child(tblPr, "w:tblLayout", {"w:type": "fixed"})
        _clear_borders(tblPr)

        for ci, cell_blocks in enumerate(cells_data):
            col_w_twips = int(col_w_pts[ci] * 20)
            word_cell = table.cell(0, ci)
            tc = word_cell._tc
            tcPr = _get_or_add(tc, "w:tcPr")

            _set_child(tcPr, "w:tcW", {"w:w": str(col_w_twips), "w:type": "dxa"})

            tcMar = _get_or_add(tcPr, "w:tcMar")
            for side in ("top", "left", "bottom", "right"):
                _set_child(tcMar, f"w:{side}", {"w:w": "0", "w:type": "dxa"})

            for p in word_cell.paragraphs:
                p._element.getparent().remove(p._element)

            cell_bx1 = boundaries[ci]
            cell_bx2 = boundaries[ci + 1]

            for block in cell_blocks:
                b = dict(block)
                b["alignment"] = _cell_alignment(block, cell_bx1, cell_bx2)
                b["space_after_pt"] = block.get("space_after_pt", 4.0)
                self._render_rich_in(b, word_cell)

    def _render_rich_in(self, block: dict, cell):
        alignment = block.get("alignment", "left")
        has_bullet = block.get("has_bullet", False)
        indent_left = block.get("indent_left_pt", 0.0)
        indent_hanging = block.get("indent_hanging_pt", 0.0)
        indent_first = block.get("indent_first_line_pt", 0.0)
        space_after = block.get("space_after_pt", 4.0)
        block_lines = block.get("block_lines") or []
        wd_align = _ALIGN_MAP.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

        if len(block_lines) > 1:
            for li, line_spans in enumerate(block_lines):
                is_last = li == len(block_lines) - 1
                para = cell.add_paragraph()
                pf = para.paragraph_format
                pf.alignment = wd_align
                pf.space_before = Pt(0)
                pf.space_after = Pt(space_after if is_last else 2.0)
                _apply_spans(para, line_spans)
        else:
            spans = block.get("block_spans", [])
            para = cell.add_paragraph()
            pf = para.paragraph_format
            pf.alignment = wd_align
            pf.space_before = Pt(0)
            pf.space_after = Pt(space_after)

            if has_bullet:
                pf.left_indent = Pt(indent_left)
                pf.first_line_indent = Pt(-indent_hanging)
                self._set_tab_stop(para, indent_left)
                body_font = spans[0]["font"] if spans else "Times New Roman"
                body_size = spans[0]["size"] if spans else 12
                br = para.add_run("•\t")
                br.font.name = body_font
                br.font.size = Pt(body_size)
            elif indent_first > 0:
                pf.first_line_indent = Pt(indent_first)

            _apply_spans(para, spans)

    def _render_hline(self, item: dict):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _render_plain(self, block: dict, page_image: Optional[Image.Image]):
        label = block.get("block_label", "text").lower()
        content = block.get("block_content", "").strip()
        bbox = block.get("block_bbox", [])
        alignment = block.get("block_alignment", "left")
        is_bold = block.get("block_bold", False)
        font = block.get("block_font", "Times New Roman")
        size = block.get("block_size", 13)

        if label in _HEADING_LABELS:
            if content:
                p = self.doc.add_heading(content, level=_HEADING_LABELS[label])
                p.alignment = _ALIGN_MAP.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)
                for run in p.runs:
                    run.font.name = font
                    run.font.size = Pt(size)
                    run.bold = is_bold
        elif label == "section_title":
            if content:
                self._plain_para(content, alignment, is_bold, font, size)
        elif label == "table":
            rows = html_to_table(content)
            if rows:
                self._plain_table(rows)
            elif content:
                self._plain_para(content, alignment, is_bold, font, size)
        elif label in _FIGURE_LABELS:
            figure_bytes = block.get("figure_bytes")
            if figure_bytes:
                try:
                    self._plain_image(Image.open(io.BytesIO(figure_bytes)))
                except Exception:
                    pass
            elif bbox and page_image:
                try:
                    self._plain_image(crop_image(page_image, bbox))
                except Exception:
                    pass
        elif label in {"list", "list_item"}:
            if content:
                self._plain_bullet(content, font, size)
        elif label == "formula":
            if content:
                self._plain_formula(content)
        else:
            if content:
                self._plain_para(content, alignment, is_bold, font, size)

    def _plain_para(self, text: str, alignment: str = "left", bold: bool = False,
                    font: str = "Times New Roman", size: int = 13):
        para = self.doc.add_paragraph()
        para.style = self.doc.styles["Normal"]
        para.paragraph_format.alignment = _ALIGN_MAP.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        run = para.add_run(text)
        run.font.name = font
        run.font.size = Pt(size)
        run.bold = bold

    def _plain_bullet(self, text: str, font: str = "Times New Roman", size: int = 13):
        para = self.doc.add_paragraph(style="List Bullet")
        run = para.add_run(text)
        run.font.name = font
        run.font.size = Pt(size)

    def _plain_formula(self, text: str):
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(11)

    def _plain_image(self, image: Image.Image, max_width: float = 6.0):
        buf = io.BytesIO(pil_to_bytes(image))
        self.doc.add_picture(buf, width=Inches(max_width))

    def _plain_table(self, rows: List[List[str]]):
        if not rows:
            return
        num_cols = max(len(r) for r in rows)
        if num_cols == 0:
            return
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                if c_idx < num_cols:
                    cell = table.cell(r_idx, c_idx)
                    cell.text = cell_text
                    if r_idx == 0:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.bold = True

    @staticmethod
    def _set_tab_stop(para, pos_pt: float):
        pPr = para._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "left")
        tab.set(qn("w:pos"), str(int(pos_pt * 20)))
        tabs.append(tab)
        pPr.append(tabs)


def _apply_spans(para, spans: List[dict]):
    for span in spans:
        run = para.add_run(span["text"])
        run.font.name = span.get("font", "Times New Roman")
        run.font.size = Pt(span["size"])
        run.bold = span.get("bold", False)
        run.italic = span.get("italic", False)
        color = span.get("color", 0)
        if color:
            r = (color >> 16) & 0xFF
            g = (color >> 8) & 0xFF
            b = color & 0xFF
            if (r, g, b) != (0, 0, 0):
                run.font.color.rgb = RGBColor(r, g, b)


def _cell_alignment(block: dict, cell_x1: float, cell_x2: float) -> str:
    bbox = block.get("block_bbox", [])
    if len(bbox) < 4:
        return block.get("alignment", "left")
    bx1, bx2 = bbox[0], bbox[2]
    bcx = (bx1 + bx2) / 2.0
    cell_cx = (cell_x1 + cell_x2) / 2.0
    cell_w = max(1.0, cell_x2 - cell_x1)

    if abs(bcx - cell_cx) / cell_w < 0.10:
        return "center"
    if (bx2 >= cell_x2 - cell_w * 0.12) and (bx1 > cell_cx):
        return "right"
    return block.get("alignment", "justify")


def _get_or_add(parent, tag: str):
    ns_tag = qn(tag)
    el = parent.find(ns_tag)
    if el is None:
        el = OxmlElement(tag)
        parent.append(el)
    return el


def _set_child(parent, tag: str, attrs: dict):
    ns_tag = qn(tag)
    existing = parent.find(ns_tag)
    if existing is not None:
        parent.remove(existing)
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k), v)
    parent.append(el)
    return el


def _clear_borders(tblPr):
    ns_tag = qn("w:tblBorders")
    existing = tblPr.find(ns_tag)
    if existing is not None:
        tblPr.remove(existing)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBdr.append(el)
    tblPr.append(tblBdr)
